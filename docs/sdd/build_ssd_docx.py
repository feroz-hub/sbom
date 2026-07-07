from __future__ import annotations

import re
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
SDD_MD = ROOT / "docs" / "sdd" / "SBOM_Analyzer_SDD.md"
OUT_DIR = ROOT / "docs" / "sdd"
SSD_MD = OUT_DIR / "SBOM_Analyzer_SSD.md"
SSD_DOCX = OUT_DIR / "SBOM_Analyzer_SSD.docx"
TEMPLATE_DOCX = OUT_DIR / "Reusable_SSD_Template.docx"


ACCENT = RGBColor(46, 116, 181)
DARK_ACCENT = RGBColor(31, 77, 120)
INK = RGBColor(31, 31, 31)
MUTED = RGBColor(94, 94, 94)
GRID = "D9E2F3"
FILL = "F2F4F7"


@dataclass
class SectionSlice:
    start: int
    end: int


def slug_heading(line: str) -> tuple[int, str] | None:
    match = re.match(r"^(#{1,6})\s+(.+?)\s*$", line)
    if not match:
        return None
    return len(match.group(1)), match.group(2).strip()


def load_sections(markdown: str) -> dict[str, SectionSlice]:
    lines = markdown.splitlines()
    headings: list[tuple[int, str, int]] = []
    for idx, line in enumerate(lines):
        parsed = slug_heading(line)
        if parsed:
            level, title = parsed
            headings.append((level, title, idx))

    sections: dict[str, SectionSlice] = {}
    for pos, (level, title, start) in enumerate(headings):
        end = len(lines)
        for next_level, _next_title, next_start in headings[pos + 1 :]:
            if next_level <= level:
                end = next_start
                break
        sections[title] = SectionSlice(start, end)
    return sections


def slice_markdown(markdown: str, title: str, heading: str | None = None) -> str:
    sections = load_sections(markdown)
    if title not in sections:
        return f"## {heading or title}\n\n`TBD: source section not found: {title}`\n"
    lines = markdown.splitlines()
    item = sections[title]
    chunk = lines[item.start : item.end]
    if heading is not None and chunk:
        parsed = slug_heading(chunk[0])
        if parsed:
            level, _ = parsed
            chunk[0] = f"{'#' * level} {heading}"
    return "\n".join(chunk).strip() + "\n"


def adjust_heading_levels(markdown: str, delta: int) -> str:
    output: list[str] = []
    for line in markdown.splitlines():
        parsed = slug_heading(line)
        if parsed:
            level, title = parsed
            level = max(1, min(6, level + delta))
            output.append(f"{'#' * level} {title}")
        else:
            output.append(line)
    return "\n".join(output)


def build_ssd_markdown() -> str:
    source = SDD_MD.read_text(encoding="utf-8")
    intro = source.split("# 1. Introduction", 1)[0].strip()

    architecture_parts = [
        slice_markdown(source, "5.1 Software System Architecture"),
        adjust_heading_levels(slice_markdown(source, "5.2 System Integration Architecture"), 1),
        adjust_heading_levels(slice_markdown(source, "5.3 Domain Terms"), 1),
    ]
    architecture = "\n\n".join(architecture_parts)
    architecture = architecture.replace("# 5.1 Software System Architecture", "## 5.1 Software System Architecture")

    module_description = slice_markdown(source, "5.4 Module Description", "5.2 Module Description").replace(
        "## 5.2 Module Description", "## 5.2 Module Description"
    )
    segregation = slice_markdown(source, "5.5 Segregation of Software", "5.3 Segregation of Software")
    boot = slice_markdown(source, "5.6 Startup and Shutdown Handling", "5.4 Bootup and Shutdown Handling")
    hosting = adjust_heading_levels(slice_markdown(source, "5.7 Background Processing Hosting"), 1)
    init_cfg = adjust_heading_levels(slice_markdown(source, "5.8 Initialization / Configuration Parameters"), 1)
    integration = adjust_heading_levels(slice_markdown(source, "5.9 Integration Approach"), 1)
    flows = adjust_heading_levels(slice_markdown(source, "5.10 Analysis Pipeline Flows"), 1)
    resource_constraints = adjust_heading_levels(slice_markdown(source, "5.11 Resource Constraints"), 1)
    dependencies = adjust_heading_levels(slice_markdown(source, "5.12 Software Dependencies, Assumptions and Constraints"), 1)

    exact = [
        intro.replace("Software Design Document For SBOM Analyzer", "Software Design Document (SSD) For SBOM Analyzer"),
        "# 1. Introduction",
        slice_markdown(source, "1.1 Purpose"),
        slice_markdown(source, "1.2 Scope"),
        slice_markdown(source, "1.3 Intended Audience"),
        slice_markdown(source, "2. Reference Documents", "2. Reference Documents"),
        slice_markdown(source, "3. Functional Description of the System"),
        slice_markdown(source, "4. Alternate Design Method Considered", "4. Alternate Design Methods Considered"),
        "# 5. Software Architecture Description",
        architecture,
        module_description,
        segregation,
        boot,
        hosting,
        init_cfg,
        integration,
        flows,
        resource_constraints,
        dependencies,
        slice_markdown(source, "6. OS Features"),
        slice_markdown(source, "7. Hardware & Software Platforms"),
        slice_markdown(source, "8. Interface Description"),
        slice_markdown(source, "9. Security Features and Threat Mitigation"),
        slice_markdown(source, "10. Individual Functional Block Explanation"),
        slice_markdown(source, "11. Data Management Features"),
        slice_markdown(source, "12. Interoperability Features"),
        slice_markdown(source, "13. Maintenance Features"),
        slice_markdown(source, "14. Security Features"),
        slice_markdown(source, "15. Product Support"),
        slice_markdown(source, "16. Standards"),
        slice_markdown(source, "17. Appendix"),
        "\n## 17.9 Companion Document Register\n\n"
        "| Companion File | Purpose |\n"
        "|---|---|\n"
        "| `SBOM_Analyzer_API_Inventory.md` / `.docx` | Complete API inventory for the functional blocks in Section 10. |\n"
        "| `SBOM_Analyzer_Data_Dictionary.md` / `.docx` | Full model, field, migration and relationship reference for Section 11. |\n"
        "| `SBOM_Analyzer_Configuration_Reference.md` / `.docx` | Configuration and feature flag source for Sections 5, 7, 13 and Appendix 17.6. |\n"
        "| `SBOM_Analyzer_Implementation_Traceability.md` / `.docx` | Source-file traceability for architecture, API, data, security and maintenance claims. |\n"
        "| `SBOM_Analyzer_Test_Traceability.md` / `.docx` | Test-to-feature coverage reference for verification planning. |\n"
        "| `SBOM_Analyzer_Open_Questions.md` / `.docx` | Design gaps, stakeholder questions and document placeholders. |\n",
    ]
    result = "\n\n".join(part.strip() for part in exact if part.strip()) + "\n"
    heading_fixes = {
        "## 2. Reference Documents": "# 2. Reference Documents",
        "### 5.1 Software System Architecture": "## 5.1 Software System Architecture",
        "### 5.2 System Integration Architecture": "### 5.1.1 System Integration Architecture",
        "### 5.3 Domain Terms": "### 5.1.2 Domain Terms",
        "### 5.5.1 API Layer": "### 5.3.1 API Layer",
        "### 5.5.2 Service Layer": "### 5.3.2 Service Layer",
        "### 5.5.3 Adapter Layer": "### 5.3.3 Adapter Layer",
        "### 5.5.4 Persistence Layer": "### 5.3.4 Persistence Layer",
        "### 5.5.5 Task Layer": "### 5.3.5 Task Layer",
        "### 5.5.6 Presentation Layer": "### 5.3.6 Presentation Layer",
        "### 5.7 Background Processing Hosting": "### 5.4.1 Background Processing Hosting",
        "### 5.8 Initialization / Configuration Parameters": "### 5.4.2 Initialization / Configuration Parameters",
        "### 5.9 Integration Approach": "### 5.4.3 Integration Approach",
        "### 5.10 Analysis Pipeline Flows": "### 5.4.4 Analysis Pipeline Flows",
        "### 5.11 Resource Constraints": "### 5.4.5 Resource Constraints",
        "### 5.12 Software Dependencies, Assumptions and Constraints": "### 5.4.6 Software Dependencies, Assumptions and Constraints",
    }
    for old, new in heading_fixes.items():
        result = result.replace(old, new)
    result = result.replace("SBOM-DOC-002", "SBOM-SSD-002")
    result = result.replace("Initial design document", "Initial SSD document")
    return result


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="C9D3E1", size="4") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_tbl_width(table, width_dxa=9360, indent_dxa=120) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")


def set_paragraph_border_bottom(paragraph, color="2E74B5", size="8") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)


def configure_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    for side in ("top_margin", "right_margin", "bottom_margin", "left_margin"):
        setattr(section, side, Inches(1))
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, ACCENT, 16, 8),
        ("Heading 2", 13, ACCENT, 12, 6),
        ("Heading 3", 12, DARK_ACCENT, 8, 4),
        ("Heading 4", 11, DARK_ACCENT, 6, 3),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    if "Code Block" not in styles:
        code = styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
    else:
        code = styles["Code Block"]
    code.font.name = "Courier New"
    code._element.rPr.rFonts.set(qn("w:ascii"), "Courier New")
    code._element.rPr.rFonts.set(qn("w:hAnsi"), "Courier New")
    code.font.size = Pt(8.5)
    code.paragraph_format.left_indent = Inches(0.15)
    code.paragraph_format.right_indent = Inches(0.15)
    code.paragraph_format.space_after = Pt(4)
    code.paragraph_format.line_spacing = 1.0


def add_run_with_inline_code(paragraph, text: str, bold=False) -> None:
    parts = re.split(r"(`[^`]+`|\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Courier New"
            run._element.rPr.rFonts.set(qn("w:ascii"), "Courier New")
            run._element.rPr.rFonts.set(qn("w:hAnsi"), "Courier New")
            run.font.size = Pt(9.5)
        elif part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            run = paragraph.add_run(part)
            run.bold = bold


def add_markdown_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    col_count = max(len(row) for row in rows)
    clean_rows = [row + [""] * (col_count - len(row)) for row in rows]
    table = doc.add_table(rows=len(clean_rows), cols=col_count)
    table.autofit = False
    set_tbl_width(table)
    set_table_borders(table)

    widths = []
    if col_count == 2:
        widths = [Inches(1.85), Inches(4.65)]
    elif col_count == 3:
        widths = [Inches(1.35), Inches(2.55), Inches(2.35)]
    elif col_count == 4:
        widths = [Inches(1.2), Inches(1.55), Inches(1.45), Inches(2.1)]
    elif col_count == 5:
        widths = [Inches(1.0), Inches(1.8), Inches(1.0), Inches(0.9), Inches(1.8)]
    else:
        widths = [Inches(6.5 / col_count)] * col_count

    for idx, width in enumerate(widths):
        for cell in table.columns[idx].cells:
            cell.width = width

    for row_idx, row in enumerate(clean_rows):
        for col_idx, text in enumerate(row):
            cell = table.cell(row_idx, col_idx)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if row_idx == 0:
                set_cell_shading(cell, FILL)
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(2)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            add_run_with_inline_code(paragraph, text.strip(), bold=row_idx == 0)
    doc.add_paragraph()


def parse_pipe_row(line: str) -> list[str]:
    stripped = line.strip().strip("|")
    return [cell.strip() for cell in stripped.split("|")]


def is_separator_row(line: str) -> bool:
    cells = parse_pipe_row(line)
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell.strip()) for cell in cells)


def add_front_matter(doc: Document) -> None:
    section = doc.sections[0]
    header = section.header
    p = header.paragraphs[0]
    p.text = "SBOM Analyzer SSD"
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.runs[0].font.color.rgb = MUTED
    p.runs[0].font.size = Pt(9)
    set_paragraph_border_bottom(p, color="D9E2F3", size="4")

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run("HCL Confidential | Page ")
    field_run = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    page_text = OxmlElement("w:t")
    page_text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    field_run._r.append(fld_begin)
    field_run._r.append(instr)
    field_run._r.append(fld_sep)
    field_run._r.append(page_text)
    field_run._r.append(fld_end)
    for run in p.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = MUTED


def add_toc(doc: Document) -> None:
    p = doc.add_paragraph()
    p.style = "Heading 1"
    p.add_run("Table of Contents")
    p = doc.add_paragraph()
    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "Right-click and update field in Word to refresh page numbers."
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(placeholder)
    run._r.append(fld_end)
    doc.add_page_break()


def markdown_to_docx(markdown: str, out_path: Path, title: str, include_toc: bool = False) -> None:
    doc = Document()
    configure_styles(doc)
    add_front_matter(doc)

    first_title_written = False
    in_code = False
    code_lines: list[str] = []
    table_rows: list[list[str]] = []
    lines = markdown.splitlines()

    if include_toc:
        add_toc(doc)

    def flush_table() -> None:
        nonlocal table_rows
        if table_rows:
            add_markdown_table(doc, table_rows)
            table_rows = []

    def flush_code() -> None:
        nonlocal code_lines
        if code_lines:
            for code_line in code_lines:
                p = doc.add_paragraph(style="Code Block")
                p.add_run(code_line[:260])
            code_lines = []

    for raw_line in lines:
        line = raw_line.rstrip()
        if line.startswith("```"):
            if in_code:
                flush_code()
                in_code = False
            else:
                flush_table()
                in_code = True
                code_lines = []
            continue
        if in_code:
            code_lines.append(line)
            continue
        if line.strip().startswith("|") and "|" in line.strip()[1:]:
            if is_separator_row(line):
                continue
            table_rows.append(parse_pipe_row(line))
            continue
        flush_table()
        if not line.strip():
            continue

        heading = slug_heading(line)
        if heading:
            level, text = heading
            if level == 1 and not first_title_written:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.space_before = Pt(10)
                p.paragraph_format.space_after = Pt(4)
                run = p.add_run(text)
                run.font.size = Pt(23)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)
                set_paragraph_border_bottom(p)
                first_title_written = True
            else:
                style_level = min(max(level, 1), 4)
                p = doc.add_paragraph(style=f"Heading {style_level}")
                add_run_with_inline_code(p, text, bold=True)
            continue

        quote = line.strip().startswith(">")
        bullet = re.match(r"^(\s*)[-*]\s+(.+)$", line)
        numbered = re.match(r"^(\s*)\d+\.\s+(.+)$", line)
        if quote:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.22)
            p.paragraph_format.right_indent = Inches(0.12)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(8)
            add_run_with_inline_code(p, line.strip().lstrip(">").strip())
        elif bullet:
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(4)
            add_run_with_inline_code(p, bullet.group(2))
        elif numbered:
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.space_after = Pt(4)
            add_run_with_inline_code(p, numbered.group(2))
        elif line.strip() == "---":
            p = doc.add_paragraph()
            set_paragraph_border_bottom(p, color=GRID, size="4")
        else:
            p = doc.add_paragraph()
            add_run_with_inline_code(p, line.strip())

    flush_table()
    flush_code()
    doc.save(out_path)


def build_template_markdown() -> str:
    return """# Reusable Software Design Document (SSD) Template

| | |
|---|---|
| **Project Code** | `<TBD>` |
| **System** | `<System name and short description>` |
| **Doc ID** | `<TBD>` |
| **Release Date** | `<TBD>` |
| **Revision** | `0.1` |
| **Status** | `Draft — for review` |

## Review / Approval

| Name | Role / Function | Signature | Date |
|---|---|---|---|
| `<TBD>` | Reviewed By — Software Architect | | |
| `<TBD>` | Reviewed By — Verification and Validation Lead | | |
| `<TBD>` | Reviewed By — Cybersecurity Reviewer | | |
| `<TBD>` | Approved By — Engineering Manager | | |
| `<TBD>` | Approved By — Quality / Regulatory | | |
| `<TBD>` | Approved By — Product Owner | | |

# 1. Introduction

## 1.1 Purpose

State why this SSD exists, what product/release it covers, and what architectural/design baseline it represents.

## 1.2 Scope

List in-scope and out-of-scope capabilities, integrations, deployment assumptions, and explicit exclusions.

## 1.3 Intended Audience

Identify the engineering, verification, cybersecurity, quality/regulatory, operations, customer, and support readers.

# 2. Reference Documents

| SN | Document ID / File Name | Version No. | Document Description |
|---|---|---|---|
| 1 | `<TBD>` | `<TBD>` | Requirement specification / PRD |
| 2 | `<TBD>` | `<TBD>` | Architecture decision record or regulatory reference |

# 3. Functional Description of the System

Describe what the system does at a high level, the main users/actors, primary workflows, and the major system outputs.

# 4. Alternate Design Methods Considered

Document rejected approaches and why they were not selected. Include constraints, trade-offs, and unresolved risks.

# 5. Software Architecture Description

## 5.1 Software System Architecture

Describe the full system architecture, runtime topology, key dependencies, protocols, and domain terms. Add system and integration diagrams.

## 5.2 Module Description

List the major modules, their responsibilities, primary files/packages, and ownership boundaries.

## 5.3 Segregation of Software

Describe the layers, boundaries, and responsibility separation, for example domain, domain services, service interfaces, application services, infrastructure, presentation, and task layers.

## 5.4 Bootup and Shutdown Handling

Describe startup checks, seed data, migrations/schema verification, background workers, scheduler startup, shutdown cleanup, and failure behavior.

# 6. OS Features

Document OS-level features used by the product, including filesystem, environment variables, networking, certificates, process model, and scheduler/worker requirements.

# 7. Hardware & Software Platforms

## 7.1 Hardware

State expected deployment hardware or minimum sizing assumptions.

## 7.2 Software

State required runtimes, databases, browsers, services, frameworks, and deployment components.

## 7.3 Software of Unknown Provenance (SOUP)

List third-party software, libraries, frameworks, and external services with version/license/source where known.

# 8. Interface Description

## 8.1 Internal Interface Features

Describe internal APIs, database contracts, service interfaces, background queues, events, and internal file/storage interfaces.

## 8.2 External Interface Features

Describe third-party APIs, identity providers, feeds, export formats, import formats, and integration error behavior.

## 8.3 User Interface Features

Describe key screens, navigation, user workflows, roles/permissions, and state/error presentation.

# 9. Security Features and Threat Mitigation

Describe threat model, authentication, authorization, tenancy, audit, input validation, secrets, transport security, data protection, and abuse controls.

# 10. Individual Functional Block Explanation

## 10.1 Global Description

### 10.1.1 Audit Trail

Describe audit events common to functional blocks.

### 10.1.2 Logging

Describe logging, correlation IDs, diagnostics, and operational visibility.

### 10.1.3 Notification

Describe user-facing notifications, async progress, alerts, emails, or explicit non-implementation.

### 10.1.4 Error Handling

Describe common validation errors, exception envelope, retry behavior, and failure recovery.

## 10.2 Functional Block Pattern

For each functional area, repeat:

| Field | Details |
|---|---|
| Global Description | `<What this block owns>` |
| Create / Register | `<Create workflow, endpoint, permissions, validation>` |
| Read / Verify | `<View/search/verify workflow>` |
| Update | `<Update workflow>` |
| Delete / Deactivate | `<Deletion or retention behavior>` |
| Audit / Logging | `<Events and diagnostics>` |
| Errors | `<Expected failures and mitigations>` |

# 11. Data Management Features

Describe storage model, entities, migrations, retention, caching, archival, backup/restore, soft delete, tenant isolation, and data quality rules.

# 12. Interoperability Features

Describe supported standards, file formats, APIs, import/export behavior, identifier mappings, and compatibility limits.

# 13. Maintenance Features

Describe observability, health checks, admin operations, scheduled jobs, cache refresh, backup, upgrade, and diagnostic runbooks.

# 14. Security Features

Describe module-level security behavior that complements Section 9: permissions, row-level scoping, secrets, secure defaults, and testing.

# 15. Product Support

Describe support boundaries, troubleshooting, logs, configuration support, deployment support, known limitations, and escalation data.

# 16. Standards

## 16.1 Standards Reference

List applicable regulatory, industry, security, API, interoperability, and coding standards.

## 16.2 Specifications Supported

List supported protocol, file-format, scoring, schema, and interface specifications.

# 17. Appendix

Include glossary, domain terms, state machines, feature flags, error codes, data maps, traceability tables, and open questions.
"""


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    ssd_md = build_ssd_markdown()
    SSD_MD.write_text(ssd_md, encoding="utf-8")
    markdown_to_docx(ssd_md, SSD_DOCX, "SBOM Analyzer SSD")
    markdown_to_docx(build_template_markdown(), TEMPLATE_DOCX, "Reusable SSD Template")
    print(SSD_DOCX)
    print(TEMPLATE_DOCX)


if __name__ == "__main__":
    main()
